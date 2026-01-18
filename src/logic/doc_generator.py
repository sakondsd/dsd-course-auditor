from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def create_verification_report(markdown_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Sarabun' # หรือ Tahoma
    style.font.size = Pt(12)

    for line in markdown_text.split('\n'):
        line = line.strip()
        if not line: continue

        if line.startswith('# '):
            p = doc.add_heading(line.replace('# ', ''), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            p = doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('|'):
            clean_line = line.replace('|', '  ').strip()
            p = doc.add_paragraph(clean_line)
            p.paragraph_format.left_indent = Pt(20)
            if "❌" in line:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer